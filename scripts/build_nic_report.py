"""Generate NORAD brand intelligence report for Nic and Jet Fuel (DOCX).

Clean, professional, restrained. Navy + single orange accent.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# -- palette
INK       = RGBColor(0x1A, 0x1A, 0x1A)   # body
NAVY      = RGBColor(0x0B, 0x2A, 0x4A)   # headings
ACCENT    = RGBColor(0xC8, 0x5A, 0x1A)   # single orange accent
MUTED     = RGBColor(0x66, 0x6A, 0x70)   # captions
SOFT      = RGBColor(0x9A, 0x9F, 0xA5)   # micro
RULE      = "C85A1A"                      # hex for borders
TINT      = "F5F2EE"                      # warm light tint for KV / header rows

doc = Document()

# -- page
for s in doc.sections:
    s.top_margin    = Cm(2.0)
    s.bottom_margin = Cm(2.0)
    s.left_margin   = Cm(2.2)
    s.right_margin  = Cm(2.2)

# -- base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = INK


# ---------- low-level helpers
def shade(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)


def border_bottom(paragraph, color="C85A1A", size="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), size)
    b.set(qn("w:space"), "2")
    b.set(qn("w:color"), color)
    pBdr.append(b)
    pPr.append(pBdr)


def cell_borders(cell, color="DDDDDD"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), color)
        tcBorders.append(e)
    tcPr.append(tcBorders)


def write(cell, text, *, bold=False, size=9.5, color=INK, italic=False):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color


# ---------- high-level helpers
def eyebrow(text, color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = color
    # letter-spacing via run properties
    rPr = r._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "40")
    rPr.append(spacing)
    return p


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    border_bottom(p, color=RULE, size="6")
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    return p


def body(text, *, italic=False, color=INK, size=10.5, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def caption(text):
    return body(text, italic=True, color=MUTED, size=9, space_after=6)


def bullets(items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if p.runs:
            p.runs[0].text = ""
        r = p.add_run(it)
        r.font.size = Pt(size)
        r.font.color.rgb = INK


def kv(rows, label_w=Cm(4.6), value_w=Cm(12.4)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (k, v) in enumerate(rows):
        c1, c2 = t.rows[i].cells
        c1.width, c2.width = label_w, value_w
        shade(c1, TINT)
        cell_borders(c1, "EAE5DE")
        cell_borders(c2, "EAE5DE")
        write(c1, k, bold=True, color=NAVY)
        write(c2, v)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    return t


def data_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in t.rows:
                r.cells[i].width = w
    # header row
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, TINT)
        cell_borders(c, "EAE5DE")
        write(c, h, bold=True, color=NAVY, size=9)
    # data rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            cell_borders(c, "EAE5DE")
            write(c, val, size=9.5)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    return t


# =====================================================================
# COVER
# =====================================================================
eyebrow("NORAD  ·  Brand Intelligence Report  ·  Vol. 01")

p = doc.add_paragraph()
r = p.add_run("Nic and Jet Fuel")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run("Nicotine-Infused Energy Drink   ·   Fort Walton Beach, FL   ·   Founded 2025")
r.font.size = Pt(10.5)
r.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run("Surfaced via Trend Hunter, October 1 2025   ·   Researched by Parallel AI + Exa Deep Research")
r.font.size = Pt(9)
r.font.color.rgb = SOFT
border_bottom(p, color=RULE, size="8")

# =====================================================================
# 1. AT A GLANCE
# =====================================================================
eyebrow("Section 01")
h1("At a Glance")
body(
    "Trend Hunter surfaced a previously-unknown brand in the emerging nicotine "
    "beverage category. NORAD ran the brand through two deep-research tools in "
    "parallel and assembled a complete BD profile in under three minutes of "
    "research time."
)
kv([
    ("Brand",           "Nic and Jet Fuel"),
    ("Legal entity",    "Beach Weekend Management LLC"),
    ("Category",        "Nicotine-infused functional energy drink"),
    ("Stage",           "Early-stage, founder-funded (no disclosed rounds)"),
    ("Headquarters",    "Fort Walton Beach, Florida"),
    ("Trigger source",  "Trend Hunter — Grace Mahas, Oct 1 2025"),
])

# =====================================================================
# 2. SOURCE TRIGGER
# =====================================================================
eyebrow("Section 02")
h1("Source Trigger")
body(
    "Discovery began with a single Trend Hunter article. Everything below the "
    "line — founders, legal entity, ingredients, retailers, competitors, "
    "digital footprint — was unknown from the article alone and was assembled "
    "by the research pipeline."
)
kv([
    ("Source",          "trendhunter.com / trends / nic-and-jet-fuel"),
    ("Published",       "October 1, 2025"),
    ("Author",          "Grace Mahas"),
    ("Headline angle",  "Clean-fuel alternative for focus"),
    ("Signal type",     "Emerging brand · Novel CPG category · Pre-funding"),
])

h2("What the article gave us")
bullets([
    "Brand name — Nic and Jet Fuel",
    "Product — Nicotina Energy, canned energy drink with microdosed nicotine",
    "Strengths — Lite (3 mg nicotine), Max (6 mg nicotine)",
    "Flavors — Citrus Surge, Grapefruit Spark",
    "Positioning — nicotine + caffeine + adaptogens + nootropics",
    "Category claim — 'first ready-to-drink nicotine energy drink'",
])

# =====================================================================
# 3. RESEARCH PIPELINE
# =====================================================================
eyebrow("Section 03")
h1("Research Pipeline")
body(
    "Two deep-research tools were run in parallel on identical brand inputs, "
    "with disambiguation prompts to filter out aviation 'jet fuel' noise. "
    "The two tools are complementary rather than competing: Exa returns "
    "clean structured data; Parallel returns narrative depth."
)

h2("Tool A  —  Parallel AI · Deep Research (Ultra model)")
kv([
    ("Mode",            "Deep Research — Ultra reasoning model"),
    ("Input format",    "Single prompt with embedded JSON template"),
    ("Pages read",      "579 pages"),
    ("Links considered","1,733 links"),
    ("Output",          "Strategic intelligence report, ~5,000 words"),
    ("Strength",        "Narrative synthesis, scenario analysis, recommendations"),
])

h2("Tool B  —  Exa · Deep Research (deep-reasoning)")
kv([
    ("Mode",            "Deep / deep-reasoning model"),
    ("Configuration",   "Structured outputs ON · System prompt in Text mode"),
    ("Result category", "Company · 25 results · Highlights ON"),
    ("Runtime",         "84.2 seconds"),
    ("Cost",            "$0.015 per run"),
    ("Output",          "Structured JSON profile (10 grouped fields)"),
    ("Strength",        "Discrete data extraction, source-cited evidence"),
])

h2("Combined effort")
data_table(
    ["Metric", "Parallel AI", "Exa", "Combined"],
    [
        ("Pages read",          "579",     "25 results",  "604+ documents"),
        ("Links considered",    "1,733",   "—",           "1,733+"),
        ("Runtime",             "~2 min",  "84.2 sec",    "~3 min wall-clock"),
        ("Direct cost",         "Included","$0.015",      "~$0.02 marginal"),
        ("Output shape",        "Narrative","JSON",        "Both layers"),
    ],
    col_widths=[Cm(3.6), Cm(4.5), Cm(4.5), Cm(4.4)],
)

# =====================================================================
# 4. COMPANY PROFILE
# =====================================================================
eyebrow("Section 04")
h1("Company Profile")
caption("Cross-verified findings from both tools, merged into a single source of truth.")

h2("Identity")
kv([
    ("Brand",           "Nic and Jet Fuel"),
    ("Legal entity",    "Beach Weekend Management LLC"),
    ("Tagline",         "Clean Fuel — sharper focus, calmer nerves, sustained drive without the crash"),
    ("Website",         "nicandjetfuel.com  (currently password-gated)"),
    ("Founded",         "2025"),
    ("Headquarters",    "119 Hollywood Blvd NW, Ste 206, Fort Walton Beach, FL"),
])

h2("People")
kv([
    ("Jayme Nabors",  "Co-Founder. Founder of NBI Properties (commercial real estate, "
                      "Fort Walton Beach). Launched the Beach Weekend brand in 2021. "
                      "Licensed pilot. University of Florida graduate. Co-host of "
                      "the 'Nicotine and Jetfuel' podcast."),
    ("Brett Divine",  "Co-Founder / Brand Strategist. Owner of Divine Media & Co. "
                      "(brand strategy and media agency). Licensed pilot. Co-host "
                      "of the 'Nicotine and Jetfuel' podcast."),
    ("BJ McCaslin*",  "Surfaced as Co-Founder / Master Beverage Formulator. "
                      "Co-founder of Holy! Water. Founded Coastal Spritz (sold to "
                      "Vita Coco). Co-founded Coco Cafe at Vita Coco. 20+ years "
                      "in beverage CPG."),
])
caption("* Surfaced by Exa only; not corroborated by Parallel. Treat as unverified until cross-checked.")

h2("Product Line — Nicotina Energy")
data_table(
    ["SKU", "Flavor", "Nicotine", "Caffeine", "Format"],
    [
        ("Citrus Surge Lite",     "Citrus Surge",     "3 mg", "75 mg", "12 oz can"),
        ("Citrus Surge Max",      "Citrus Surge",     "6 mg", "75 mg", "12 oz can"),
        ("Grapefruit Spark Lite", "Grapefruit Spark", "3 mg", "75 mg", "12 oz can"),
        ("Grapefruit Spark Max",  "Grapefruit Spark", "6 mg", "75 mg", "12 oz can"),
    ],
    col_widths=[Cm(4.8), Cm(4.0), Cm(2.6), Cm(2.6), Cm(3.0)],
)

h2("Ingredient stack  (per can)")
bullets([
    "Nicotine — 3 mg (Lite) or 6 mg (Max). Source not publicly disclosed.",
    "Natural caffeine — 75 mg, derived from yerba mate",
    "Ashwagandha extract (adaptogen)",
    "L-theanine (nootropic)",
    "Rhodiola rosea extract (adaptogen)",
    "Yerba mate extract",
    "Sugar-free  ·  5 cal (Citrus Surge)  ·  10 cal (Grapefruit Spark)",
])
caption("Per-ingredient dosages beyond nicotine and caffeine are not publicly disclosed.")

h2("Pricing")
kv([
    ("DTC starting price",  "$29.00 (4-pack, 12 oz cans)"),
    ("Trade pricing",       "Not publicly disclosed"),
])

h2("Distribution")
body("Channels:", color=NAVY, space_after=2)
bullets([
    "DTC via nicandjetfuel.com (Shopify) — currently password-gated",
    "Online vape and smoke-shop retailers",
    "Wholesale (wholesale page listed on brand site)",
])
body("Confirmed retailers and distributors:", color=NAVY, space_after=2)
bullets([
    "MyVPro.com  —  online vape retailer",
    "Supply Center USA  —  wholesale",
    "VassDistro  —  wholesale",
])

h2("Audience & Positioning")
kv([
    ("Target audience", "Adults 21+ already using nicotine products (pouches, vape) "
                        "seeking an alternative delivery format; performance-oriented "
                        "consumers seeking focus without sugar crash."),
    ("Positioning",     "'Clean fuel' for cognitive performance — distances from "
                        "combustibles and vaping; emphasizes microdosing, adaptogens, "
                        "and a sugar-free formulation."),
    ("Category claim",  "First ready-to-drink nicotine energy drink."),
])

h2("Competitive Landscape")
data_table(
    ["Brand", "Format", "Notes"],
    [
        ("EXIT (Anarchy Beverages)", "Beverage",    "Nicotine + caffeine shot — 6 mg nicotine, 40 mg caffeine, 250 ml"),
        ("STiMY Energy",             "Beverage",    "100 mg caffeine + 5 mg nicotine drink"),
        ("Zyn",                      "Pouch",       "Nicotine pouches, 3–6 mg, category leader"),
        ("FRE",                      "Pouch",       "Synthetic nicotine pouches, 3–12 mg"),
        ("Lucy",                     "Pouch / gum", "Nicotine pouches and gum"),
        ("Jolt",                     "Pouch",       "Nicotine pouch brand"),
        ("NicShot",                  "Beverage",    "Nicotine shot beverage"),
        ("Wild Hempettes",           "Hybrid",      "Hemp + nicotine product"),
    ],
    col_widths=[Cm(4.6), Cm(3.4), Cm(9.0)],
)

h2("Digital Footprint")
kv([
    ("Instagram",         "@nicandjetfuel  (active)"),
    ("TikTok / LinkedIn / X", "No verified handles"),
    ("Podcast",           "'Nicotine and Jetfuel' — co-hosted by Jayme Nabors and Brett Divine"),
])

# =====================================================================
# 5. VALUE ADDED
# =====================================================================
eyebrow("Section 05")
h1("Value Added by Research")
body(
    "The Trend Hunter article gave us a brand name and a product description. "
    "The research pipeline turned that into a complete BD profile in roughly "
    "three minutes."
)
data_table(
    ["Data point", "From article", "After research"],
    [
        ("Legal entity",      "—",                           "Beach Weekend Management LLC"),
        ("Founders",          "—",                           "Jayme Nabors, Brett Divine  (+ BJ McCaslin unverified)"),
        ("Headquarters",      "—",                           "Fort Walton Beach, FL — full street address"),
        ("Ingredient stack",  "Caffeine + adaptogens (vague)","Yerba-mate caffeine, ashwagandha, L-theanine, rhodiola"),
        ("Pricing",           "—",                           "$29.00 starting (DTC)"),
        ("Retailers",         "—",                           "MyVPro, Supply Center USA, VassDistro"),
        ("Competitors",       "—",                           "8 named brands across pouch + beverage formats"),
        ("Digital footprint", "—",                           "Instagram + podcast + DTC Shopify"),
    ],
    col_widths=[Cm(4.2), Cm(4.6), Cm(8.2)],
)

# =====================================================================
# 6. DATA GAPS
# =====================================================================
eyebrow("Section 06")
h1("Open Data Gaps")
caption("Items not surfaced by either tool — flagged for a second research pass.")
bullets([
    "Nicotine source  —  synthetic vs. tobacco-derived",
    "Co-packer / contract manufacturer",
    "Funding source  (no rounds disclosed; likely founder-funded)",
    "Full MSRP and trade pricing ladder",
    "BJ McCaslin's connection to the brand  —  to be verified",
    "USPTO trademark filings",
])

# =====================================================================
# 7. SOURCES
# =====================================================================
eyebrow("Section 07")
h1("Primary Sources")
bullets([
    "nicandjetfuel.com  —  brand site (DTC, password-gated)",
    "trendhunter.com/trends/nic-and-jet-fuel  —  Trend Hunter article (Grace Mahas, Oct 1 2025)",
    "myvpro.com/products/nicotina-energy-lite-drink  —  retailer listing",
    "divinemedia.co/brett-divine  —  Brett Divine bio",
    "nbiproperties.com/our-team/20405  —  Jayme Nabors bio",
    "theorg.com/org/holy-water-1/org-chart/bj-mccaslin  —  BJ McCaslin (Holy! Water context)",
    "linkedin.com  —  'Shelf Life Story' brand coverage (Aug 20 2025)",
    "medium.com/@shelflifestory  —  nicotine energy-drink history piece",
])

# =====================================================================
# FOOTER
# =====================================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
border_bottom(p, color=RULE, size="6")

p = doc.add_paragraph()
r = p.add_run("NORAD  ·  Brand intelligence engine, built by Growth Maschine.")
r.italic = True
r.font.size = Pt(8.5)
r.font.color.rgb = SOFT

out = "NIC_AND_JET_FUEL_REPORT.docx"
doc.save(out)
print(f"Saved {out}")
