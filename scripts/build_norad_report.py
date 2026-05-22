"""NORAD Brand Intelligence Report — Vol. 01.

Two companies (Nic and Jet Fuel, Ultra Pouches), two tools each
(Parallel AI Deep Research, Exa Deep Research). Clean, professional,
restrained design — navy + single warm orange accent.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── palette ───────────────────────────────────────────────────────────
INK    = RGBColor(0x1A, 0x1A, 0x1A)
NAVY   = RGBColor(0x0B, 0x2A, 0x4A)
ACCENT = RGBColor(0xC8, 0x5A, 0x1A)
MUTED  = RGBColor(0x66, 0x6A, 0x70)
SOFT   = RGBColor(0x9A, 0x9F, 0xA5)
RULE   = "C85A1A"
TINT   = "F5F2EE"
BORDER = "EAE5DE"

doc = Document()

for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(2.2)

base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10.5)
base.font.color.rgb = INK


# ── helpers ───────────────────────────────────────────────────────────
def shade(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)


def border_bottom(p, color=RULE, size="6"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), size)
    b.set(qn("w:space"), "2")
    b.set(qn("w:color"), color)
    pBdr.append(b)
    pPr.append(pBdr)


def cell_borders(cell, color=BORDER):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), color)
        tcB.append(e)
    tcPr.append(tcB)


def write(cell, text, *, bold=False, size=9.5, color=INK, italic=False):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color


def eyebrow(text, color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = color
    rPr = r._element.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), "40")
    rPr.append(sp)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY
    border_bottom(p, RULE, "6")


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = ACCENT


def body(text, *, italic=False, color=INK, size=10.5, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color


def caption(text):
    body(text, italic=True, color=MUTED, size=9, space_after=6)


def bullets(items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if p.runs:
            p.runs[0].text = ""
        r = p.add_run(it)
        r.font.size = Pt(size)
        r.font.color.rgb = INK


def kv(rows, label_w=Cm(4.4), value_w=Cm(12.6)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (k, v) in enumerate(rows):
        c1, c2 = t.rows[i].cells
        c1.width, c2.width = label_w, value_w
        shade(c1, TINT)
        cell_borders(c1)
        cell_borders(c2)
        write(c1, k, bold=True, color=NAVY)
        write(c2, v)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def data_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in t.rows:
                r.cells[i].width = w
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, TINT)
        cell_borders(c)
        write(c, h, bold=True, color=NAVY, size=9)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            cell_borders(c)
            write(c, val, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def quote_block(text):
    """Indented quote-style block for query / prompt examples."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), RULE)
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED


def page_break():
    doc.add_page_break()


# =====================================================================
# COVER
# =====================================================================
eyebrow("NORAD  ·  Brand Intelligence Report  ·  Vol. 01")
p = doc.add_paragraph()
r = p.add_run("Two Brands. Two Tools.\nFrom Article to Profile.")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run("Nic and Jet Fuel   ·   Ultra Pouches")
r.font.size = Pt(11)
r.font.color.rgb = MUTED
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run(
    "A demonstration of NORAD's discovery-to-intelligence pipeline, "
    "comparing Parallel AI Deep Research and Exa Deep Research on two "
    "brands surfaced via Trend Hunter."
)
r.font.size = Pt(9.5)
r.font.color.rgb = SOFT
border_bottom(p, RULE, "8")


# =====================================================================
# 01 — THE PIPELINE
# =====================================================================
eyebrow("Section 01")
h1("The NORAD Pipeline")
body(
    "Every brand in this report follows the same path. A Trend Hunter article "
    "is the trigger. From it, NORAD extracts a brand context and writes a "
    "research query for each tool. The same query is then sent to two deep-"
    "research engines running in parallel. The outputs are merged into a "
    "single brand profile."
)

h2("The four-step flow")
data_table(
    ["Step", "What happens", "Output"],
    [
        ("1. Trigger",   "Trend Hunter article surfaces a new brand or category",         "Brand name + product description"),
        ("2. Extract",   "NORAD extracts SKUs, claims, disambiguation anchors, signals",  "Structured brand context"),
        ("3. Query",     "A research prompt is generated for each tool with the context", "Two parallel research queries"),
        ("4. Synthesize","Both tools run; outputs are merged and gaps flagged",            "Unified brand profile"),
    ],
    col_widths=[Cm(3.0), Cm(8.5), Cm(5.5)],
)

body(
    "This report walks through that flow end-to-end for two brands. For each "
    "company, you see the exact inputs that were passed to the tools, the "
    "outputs from each tool, and the consolidated profile."
)


# =====================================================================
# 02 — NIC AND JET FUEL
# =====================================================================
page_break()
eyebrow("Section 02  ·  Company One")
h1("Nic and Jet Fuel")
body("Nicotine-infused energy drink   ·   Fort Walton Beach, FL   ·   Founded 2025",
     italic=True, color=MUTED, size=10)

# 02.1 What we passed
h2("2.1   What We Passed to the Tools")
body(
    "The Trend Hunter article gave NORAD a brand name, a product description, "
    "and a positioning angle. From that, the following inputs were generated."
)

h3("From the article (raw)")
kv([
    ("Source",         "Trend Hunter — Grace Mahas, October 1, 2025"),
    ("Brand",          "Nic and Jet Fuel"),
    ("Product",        "Nicotina Energy — nicotine-infused canned energy drink"),
    ("Strengths",      "Lite (3 mg nicotine) · Max (6 mg nicotine)"),
    ("Flavors",        "Citrus Surge · Grapefruit Spark"),
    ("Positioning",    "Sugar-free; microdosed nicotine + caffeine + adaptogens + nootropics; 'clean fuel' for focus"),
])

h3("Research query sent to both tools")
quote_block(
    '"Nic and Jet Fuel" company — nicotine-infused canned energy drink brand. '
    'SKUs: Lite 3MG nicotine, Max 6MG nicotine. Flavors: Citrus Surge, Grapefruit Spark. '
    'Sugar-free, microdosed nicotine plus caffeine, adaptogens, nootropics. Positioned '
    'as "clean fuel" for performance and focus. Surfaced via Trend Hunter October 1 2025 '
    'by Grace Mahas.\n\n'
    'Research: founders and prior companies, leadership, funding history, full ingredient '
    'stack per SKU with dosages, co-packer, retailers, US price points, social handles, '
    'recent news 12 months, competitors, regulatory status.\n\n'
    'IGNORE: aviation jet fuel, JP-8, military fuel, JetBlue, or any unrelated '
    '"jet fuel" product.'
)
caption("Disambiguation prompt prevents the tools from drifting into aviation fuel content.")

# 02.2 Parallel
h2("2.2   Tool A  —  Parallel AI Deep Research")
kv([
    ("Mode",             "Deep Research · Ultra reasoning model"),
    ("Pages read",       "579 pages"),
    ("Links considered", "1,733 links"),
    ("Output",           "Narrative intelligence report, ~5,000 words"),
])

h3("Key findings")
bullets([
    "Legal entity confirmed: Beach Weekend Management LLC",
    "Address: 119 Hollywood Blvd NW, Ste 206, Fort Walton Beach, FL",
    "Founders identified: Jayme Nabors III (named in public filings), Brett Divine (Divine Media & Co., podcast co-host)",
    "Full 4-SKU portfolio mapped with 75 mg caffeine per can",
    "Wholesale distributors named: Supply Center USA, VassDistro",
    "Competitive landscape positioned across pouch + beverage formats",
    "DTC site identified as currently password-gated",
    "Pricing data flagged as missing (no public MSRP / trade ladder)",
])
caption("Parallel's strength: narrative synthesis, scenario analysis, strategic recommendations.")

# 02.3 Exa
h2("2.3   Tool B  —  Exa Deep Research")
kv([
    ("Mode",     "Deep / deep-reasoning model"),
    ("Results",  "25 documents · Highlights ON"),
    ("Runtime",  "84.2 seconds"),
    ("Cost",     "$0.015"),
    ("Output",   "Structured JSON profile"),
])

h3("Key findings")
bullets([
    "Legal entity confirmed: Beach Weekend Management LLC (matches Parallel)",
    "Three founders surfaced — Nabors, Divine, plus BJ McCaslin (Co-Founder / Master Beverage Formulator, Holy! Water, Coastal Spritz)",
    "Full ingredient stack: 75 mg yerba-mate caffeine, ashwagandha, L-theanine, rhodiola rosea, yerba mate extract",
    "Retailer identified: MyVPro.com",
    "DTC starting price extracted: $29.00 (4-pack)",
    "8 named competitors with positioning notes (EXIT, STiMY, Zyn, FRE, Lucy, Jolt, NicShot, Wild Hempettes)",
    "Social handle: @nicandjetfuel (Instagram only)",
    "18 primary sources cited",
])
caption("Exa's strength: clean structured data, discrete fields, source-cited evidence.")

# 02.4 Consolidated
h2("2.4   Consolidated Profile")
kv([
    ("Brand",          "Nic and Jet Fuel"),
    ("Legal entity",   "Beach Weekend Management LLC"),
    ("HQ",             "Fort Walton Beach, FL"),
    ("Founded",        "2025"),
    ("Founders",       "Jayme Nabors III, Brett Divine  (BJ McCaslin surfaced by Exa, pending verification)"),
    ("Product line",   "Nicotina Energy — 4 SKUs (Citrus Surge / Grapefruit Spark × Lite 3 mg / Max 6 mg)"),
    ("Active stack",   "Nicotine 3-6 mg · 75 mg yerba-mate caffeine · ashwagandha · L-theanine · rhodiola · sugar-free"),
    ("DTC price",      "$29.00 (4-pack)"),
    ("Channels",       "DTC (password-gated) · MyVPro · Supply Center USA · VassDistro"),
    ("Social",         "@nicandjetfuel (Instagram) · 'Nicotine and Jetfuel' podcast"),
    ("Funding",        "No disclosed rounds — likely founder-funded"),
])


# =====================================================================
# 03 — ULTRA POUCHES
# =====================================================================
page_break()
eyebrow("Section 03  ·  Company Two")
h1("Ultra Pouches")
body("Nicotine-free functional pouch   ·   Brooklyn, NY   ·   Founded 2025",
     italic=True, color=MUTED, size=10)

# 03.1 What we passed
h2("3.1   What We Passed to the Tools")
body(
    "Ultra was surfaced via a Trend Hunter article tied to a $11M Series A "
    "announcement. The brand name 'Ultra' overlaps with many unrelated "
    "entities, so the query required heavy disambiguation."
)

h3("From the article (raw)")
kv([
    ("Source",        "Trend Hunter — Kalin Ned, January 21, 2026"),
    ("Brand",         "Ultra"),
    ("Product",       "Non-addictive oral pouch — no nicotine, no caffeine"),
    ("Active",        "Paraxanthine (via Enfinity®, TSI Group / Ingenious Ingredients)"),
    ("Other actives", "L-theanine · Alpha GPC · B-vitamins · ginseng"),
    ("Funding",       "$11M Series A (BusinessWire, Jan 21 2026)"),
    ("Positioning",   "Cognitive enhancement for professionals + athletes"),
])

h3("Research query sent to both tools")
quote_block(
    '"Ultra" company — brand making NON-ADDICTIVE oral pouches with NO nicotine and '
    'NO caffeine. Raised $11 million Series A announced via BusinessWire January 21 2026. '
    'Ingredients: paraxanthine delivered as Enfinity (licensed from TSI Group / Ingenious '
    'Ingredients), L-theanine, Alpha GPC, B vitamins, ginseng extract. Sugar-free. '
    'Targets professionals and athletes.\n\n'
    'Research: legal entity, founders and prior companies, leadership, $11M Series A '
    'details (lead investor, participating investors, valuation, date), full SKU list, '
    'Enfinity / paraxanthine licensing relationship, retailers, US price points, social '
    'handles, recent news 12 months, competitors, regulatory status of paraxanthine.\n\n'
    'IGNORE: Ultra music festival, Ultra Mobile MVNO, Ulta Beauty, Ultra Tune, '
    'Ultra Records, or any other "Ultra"-named company.'
)
caption("The disambiguation list is what made this query work — 'Ultra' is one of the most polluted brand namespaces online.")

# 03.2 Parallel
h2("3.2   Tool A  —  Parallel AI Deep Research")
kv([
    ("Mode",             "Deep Research · Ultra reasoning model"),
    ("Pages read",       "774 pages"),
    ("Links considered", "1,631 links"),
    ("Output",           "Narrative intelligence report, ~5,000 words"),
])

h3("Key findings")
bullets([
    "Legal entity: Ultra, Inc.",
    "Founder: Eric Drymer (previously co-founder of Studyverse)",
    "1,000,000 cans sold in first 6 months post-launch (launched May 2025)",
    "Series A $11M led by Left Lane Capital with participation from founders of Harry's, Grüns, Rockstar Energy",
    "Celebrity-athlete investors: Joe Burrow, Lindsey Vonn, Dak Prescott",
    "SKUs: Focus Pouches (Cool Mint, Wintergreen, Tropical, Watermelon) + Sleep Pouches (Honey Lemon)",
    "Paraxanthine 100 mg per pouch; 15 pouches per tin",
    "Pricing: $16 per tin · Subscribe & Save bundles at $34.20 for 3-pack",
    "Channels active: DTC (takeultra.com) + Amazon; retail expansion underway (new Head of Retail role identified)",
    "Athlete users named: Nate Diaz, Rampage Jackson",
    "TSI Group identified as exclusive global distributor for Enfinity",
])
caption("Parallel found the 1M-cans velocity story and the new Head of Retail hire — both retail-readiness signals Exa missed.")

# 03.3 Exa
h2("3.3   Tool B  —  Exa Deep Research")
kv([
    ("Mode",     "Deep / deep-reasoning model"),
    ("Results",  "10 documents (deep variant) · Highlights ON"),
    ("Runtime",  "61.1 seconds"),
    ("Cost",     "$0.015"),
    ("Output",   "Structured JSON profile"),
])

h3("Key findings")
bullets([
    "Legal entity: Ultra Pouches (fka Nuro Pouches, Inc.) — the prior name is a meaningful data point Parallel missed",
    "Headquarters: 325 Kent Ave, Brooklyn, NY 11249 — full street address",
    "Founder bio: Eric Drymer — USC + Sciences Po; prior roles at Studyverse, Founders Intelligence, Silicon Valley Bank, European Geostrategy",
    "Series A details corroborated: $11M, Jan 20 2026, Left Lane Capital lead",
    "Full Sleep Pouch formulation: melatonin 1 mg + magnesium gluconate 30 mg + lemon balm 12 mg + chamomile 10 mg + passionflower 10 mg",
    "Co-manufacturer identified: NNB Nutrition (manufactures Enfinity for TSI Group)",
    "Channels: DTC + Amazon + TikTok Shop (@ultrapouches)",
    "11 named competitors (Lucy Breaks, FRE, Rogue, Black Buffalo, Grinds, Neuro, Magic Mind, IQ Pouch, Sesh, ALP, C.R.E.A.M.)",
    "Regulatory: paraxanthine self-affirmed GRAS since Nov 2021 — NOT FDA-affirmed (important nuance)",
    "23 primary sources cited",
])
caption("Exa found the legal-name change (Nuro Pouches → Ultra) and the co-manufacturer (NNB Nutrition) — Parallel did not.")

# 03.4 Consolidated
h2("3.4   Consolidated Profile")
kv([
    ("Brand",          "Ultra Pouches"),
    ("Legal entity",   "Ultra, Inc. (fka Nuro Pouches, Inc.)"),
    ("HQ",             "325 Kent Ave, Brooklyn, NY 11249"),
    ("Founded",        "May 2025"),
    ("Founder",        "Eric Drymer (Founder & CEO) — ex-Studyverse, SVB, Founders Intelligence"),
    ("Funding",        "$11M Series A — Jan 20 2026 — led by Left Lane Capital"),
    ("Notable backers","Founders of Harry's, Grüns, Rockstar Energy · Joe Burrow · Dak Prescott · Lindsey Vonn"),
    ("Product line",   "Ultra Focus Pouches (4 flavors) · Ultra Sleep Pouches (Honey Lemon)"),
    ("Active stack",   "Enfinity® paraxanthine 100 mg · L-theanine · Alpha GPC · B6/B12 · ginseng"),
    ("Manufacturing",  "Enfinity manufactured by NNB Nutrition · TSI Group is exclusive global distributor"),
    ("Pricing",        "$16 per 15-pouch tin (~$1.07 / pouch) · Subscribe & Save bundles"),
    ("Channels",       "DTC (takeultra.com) · Amazon · TikTok Shop · retail rollout in progress"),
    ("Social",         "Instagram, TikTok, LinkedIn, X — all @ultrapouches"),
    ("Traction",       "1M cans sold in first 6 months; positioned as #1 nicotine-free pouch globally"),
    ("Regulatory",     "Paraxanthine self-affirmed GRAS (Nov 2021) — NOT FDA-affirmed"),
])


# =====================================================================
# 04 — BOTH BRANDS, SIDE BY SIDE
# =====================================================================
page_break()
eyebrow("Section 04")
h1("Both Brands at a Glance")
body(
    "Both brands were surfaced via Trend Hunter within four months of each "
    "other. Both target performance-minded adults with functional-stimulant "
    "alternatives. But they sit on opposite ends of the regulatory and "
    "commercial spectrum — and the contrast is exactly the kind of pattern "
    "NORAD is built to surface for BD partners."
)

data_table(
    ["Dimension", "Nic and Jet Fuel", "Ultra Pouches"],
    [
        ("Category",       "Nicotine-infused energy drink",      "Nicotine-free functional pouch"),
        ("Format",         "12 oz canned beverage",              "15-count oral pouch tin"),
        ("Core active",    "Nicotine 3–6 mg + caffeine 75 mg",   "Paraxanthine 100 mg (Enfinity®)"),
        ("Legal entity",   "Beach Weekend Management LLC",       "Ultra, Inc. (fka Nuro Pouches)"),
        ("HQ",             "Fort Walton Beach, FL",              "Brooklyn, NY"),
        ("Founded",        "2025",                                "May 2025"),
        ("Founders",       "Nabors, Divine (+ McCaslin TBC)",    "Eric Drymer"),
        ("Funding",        "Founder-funded (no rounds)",          "$11M Series A — Left Lane Capital"),
        ("Pricing",        "$29 / 4-pack DTC",                    "$16 / 15-pouch tin"),
        ("Channels",       "DTC (paused) · vape retail · wholesale", "DTC · Amazon · TikTok Shop · retail rollout"),
        ("Traction",       "Emerging / pre-revenue scale",        "1M cans in 6 months · #1 nicotine-free globally"),
        ("Regulatory",     "Novel — FDA territory unclear",       "Paraxanthine self-affirmed GRAS"),
        ("BD posture",     "Watch · regulatory clarity needed",   "Engage · funded · scaling"),
    ],
    col_widths=[Cm(3.4), Cm(6.8), Cm(6.8)],
)

h2("Reading the contrast")
body(
    "Nic and Jet Fuel is the high-risk, high-novelty play — a founder-funded "
    "first-mover with a clean SKU architecture but no funding signal and "
    "unclear regulatory footing. Ultra Pouches is the high-confidence, "
    "high-velocity play — institutionally backed, with celebrity-athlete "
    "social proof, a proven first-year sell-through, and an active retail "
    "expansion. A BD partner can engage Ultra immediately; Nic and Jet Fuel "
    "is a watch-list candidate pending regulatory clarification."
)


# =====================================================================
# 05 — TOOL COMPARISON
# =====================================================================
eyebrow("Section 05")
h1("Tool Comparison  —  Parallel AI vs. Exa")
body(
    "Both tools ran on identical queries for both companies. The numbers "
    "below are direct measurements taken from the tool dashboards."
)

h2("5.1   Per-run metrics")
data_table(
    ["Metric", "Nic and Jet Fuel — Parallel", "Nic and Jet Fuel — Exa",
     "Ultra — Parallel", "Ultra — Exa"],
    [
        ("Pages read / results",  "579 pages",      "25 results",   "774 pages",      "10 results"),
        ("Links considered",       "1,733",          "—",            "1,631",          "—"),
        ("Runtime",                "~2 minutes",     "84.2 seconds", "~2 minutes",     "61.1 seconds"),
        ("Direct cost",            "Plan-included",  "$0.015",       "Plan-included",  "$0.015"),
        ("Output format",          "Markdown report","JSON",         "Markdown report","JSON"),
    ],
    col_widths=[Cm(3.0), Cm(3.5), Cm(3.5), Cm(3.5), Cm(3.5)],
)

h2("5.2   Totals across both brands")
data_table(
    ["Metric", "Parallel AI", "Exa", "Combined"],
    [
        ("Documents consumed", "1,353 pages read",       "35 result documents",   "1,388+ documents"),
        ("Links considered",   "3,364 links",             "—",                     "3,364+ links"),
        ("Total runtime",      "~4 minutes",              "145.3 seconds",         "~6.5 minutes wall-clock"),
        ("Total direct cost",  "Included in plan",        "$0.030",                "~$0.03 marginal"),
        ("Output volume",      "~10,000 words narrative", "Two structured JSON",   "Narrative + structured"),
    ],
    col_widths=[Cm(3.6), Cm(4.5), Cm(4.5), Cm(4.4)],
)

h2("5.3   Where each tool wins")
data_table(
    ["Strength", "Parallel AI", "Exa"],
    [
        ("Structured data extraction",  "Good",  "Excellent — clean JSON ready for UI"),
        ("Narrative synthesis",          "Excellent — consulting-grade reports", "Limited"),
        ("Discrete data points",         "Good",  "Excellent — funding, prior-name, manufacturer"),
        ("Velocity / momentum signals",  "Excellent — found 1M-cans story",      "Weaker"),
        ("Regulatory nuance",            "Good",  "Excellent — GRAS self-affirmation caveat"),
        ("Scenario analysis & roadmaps", "Excellent",                            "Not designed for this"),
        ("Speed",                        "~2 min",                                "~60–85 sec"),
        ("Marginal cost per run",        "Plan-included",                         "$0.015"),
    ],
    col_widths=[Cm(4.4), Cm(6.3), Cm(6.3)],
)

h2("5.4   Recommendation")
body(
    "Run both, every time. Exa fills the structured data layer that drives "
    "the NORAD UI — company cards, comparison tables, exportable profiles. "
    "Parallel fills the narrative layer that drives the BD brief — context, "
    "scenarios, recommendations. Neither tool covers what the other does. "
    "The combined cost is roughly three cents per company and the combined "
    "runtime is roughly three minutes — trivial compared to the depth of "
    "the resulting profile."
)


# =====================================================================
# FOOTER
# =====================================================================
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
border_bottom(p, RULE, "6")

p = doc.add_paragraph()
r = p.add_run("NORAD  ·  Brand intelligence engine, built by Growth Maschine.")
r.italic = True
r.font.size = Pt(8.5)
r.font.color.rgb = SOFT

out = "NORAD_REPORT_VOL01.docx"
doc.save(out)
print(f"Saved {out}")
