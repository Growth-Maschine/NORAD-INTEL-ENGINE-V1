"""Generate NORAD_DISCOVERY_QUESTIONS.docx — short, clean, normal stakeholder doc.

Critical questions first, operational questions in a separate section.
No defaults, no response boxes, no checking legend. Just questions + brief why.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x2A, 0x4D)
ACCENT = RGBColor(0x2E, 0x6B, 0xB8)
GREY = RGBColor(0x55, 0x5B, 0x66)
LIGHT_GREY = RGBColor(0x88, 0x8E, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

NAVY_HEX = "1F2A4D"
ACCENT_HEX = "2E6BB8"
AMBER_HEX = "B66A00"
PALE_NAVY_HEX = "EEF2F8"


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="D8DCE3", sz="0"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single' if sz != "0" else 'nil')
        b.set(qn('w:sz'), sz if sz != "0" else "0")
        b.set(qn('w:color'), color)
        borders.append(b)
    tc_pr.append(borders)


def remove_default_para(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


def set_cell_margins(cell, top=80, left=120, bottom=80, right=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tc_pr.append(tcMar)


def add_run(p, text, *, size=10, bold=False, italic=False, color=None, name="Calibri"):
    r = p.add_run(text)
    r.font.name = name
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def cell_para(cell, *, before=0, after=0, line_spacing=1.2, alignment=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        p.alignment = alignment
    return p


def add_para(doc, text, *, before=0, after=4, size=10, bold=False, italic=False,
             color=None, alignment=None, line_spacing=1.3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        p.alignment = alignment
    add_run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_spacer(doc, points=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, " ", size=points)


# ============ DOCUMENT ============

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


# ---- HEADER STRIP (one compact band, not a full cover) ----
header = doc.add_table(rows=1, cols=1)
hcell = header.rows[0].cells[0]
hcell.width = Cm(17)
shade_cell(hcell, NAVY_HEX)
set_cell_borders(hcell, color=NAVY_HEX, sz="0")
set_cell_margins(hcell, top=120, bottom=120, left=180, right=180)
remove_default_para(hcell)
p = cell_para(hcell, after=2)
add_run(p, "PROJECT NORAD 1", size=8.5, bold=True, color=RGBColor(0xB8, 0xC4, 0xDC))
p = cell_para(hcell, after=0)
add_run(p, "Discovery Questions for BAT", size=18, bold=True, color=WHITE)


# ---- META LINE ----
add_spacer(doc, 6)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.2
add_run(p, "From  ", size=9, color=LIGHT_GREY)
add_run(p, "Growth Machine", size=9, bold=True, color=NAVY)
add_run(p, "      To  ", size=9, color=LIGHT_GREY)
add_run(p, "BAT — BD Leadership", size=9, bold=True, color=NAVY)
add_run(p, "      Date  ", size=9, color=LIGHT_GREY)
add_run(p, "April 2026", size=9, bold=True, color=NAVY)


# ---- BRIEF INTRO (one short paragraph) ----
add_spacer(doc, 10)
add_para(doc,
    "The questions below are the decisions we need from BAT before engineering "
    "kickoff. They are split into two groups: critical decisions that define "
    "what NORAD is, and operational decisions that define how it runs.",
    size=10.5, color=GREY, after=10, line_spacing=1.4)


# ---- QUESTIONS DATA ----
PART_A = [  # Critical — define the product
    ("Q1", "PRODUCT SHAPE",
     "Should NORAD v1 ship as a scheduled-batch dashboard, or include interactive natural-language search?",
     "This is the largest scope decision. Interactive search is a different product shape, not an add-on."),

    ("Q2", "VOLUME",
     "What is the target number of qualified opportunities surfaced per week?",
     "Volume drives source-list breadth and the size of the BAT review queue."),

    ("Q3", "VERTICAL SCOPE",
     "What are the in-scope and out-of-scope verticals for v1?",
     "Each vertical added requires source curation and rubric tuning. Wider fence = noisier signal."),

    ("Q4", "SCORING",
     "Are the five scoring metrics and their weights from the brief locked, or open to refinement after the first 50 real outputs?",
     "Most rubrics shift after operators see real outputs. Treating v1 as a hypothesis is more accurate than treating it as final."),

    ("Q5", "SCORING",
     "What minimum total score should an opportunity reach before entering the BAT review queue?",
     "The threshold is the throttle on review fatigue. Too low and analysts drown; too high and the queue is empty."),

    ("Q6", "RED FLAGS",
     "Are there BAT-specific red-flag rules beyond what is in the brief — for example ESG exclusions, related-party patterns, or prior-relationship history?",
     "BAT-specific rules must be captured before screening logic is locked, otherwise false positives reach the review queue."),

    ("Q7", "QUALITY",
     "Will BAT pre-load approximately 50 curated example companies (good fits, bad fits, red-flag examples) before launch?",
     "Examples calibrate the scoring engine from day one and avoid a 4–6 week post-launch tuning period."),
]

PART_B = [  # Operational — outreach, integrations, launch
    ("Q8", "APPROVAL",
     "What is the approval chain for outbound messages — BD lead alone, or BD + Marketing + Legal?",
     "The approval chain directly sets outbound velocity and the complexity of the approval interface."),

    ("Q9", "SENDER IDENTITY",
     "Should outbound emails come from a corporate alias, or from an individual BD analyst's mailbox?",
     "Sender identity drives reply rates and determines whether per-user OAuth is required."),

    ("Q10", "CRM",
     "Confirm HubSpot is the CRM. Which pipeline and deal stages should NORAD write into?",
     "Writing into the wrong pipeline pollutes BAT's existing CRM analytics and confuses BD operators."),

    ("Q11", "AUTHENTICATION",
     "What identity provider should NORAD authenticate against — Microsoft Entra (Azure AD), Okta, Google Workspace, or standard email?",
     "Determines the SSO integration path and the configuration effort on BAT's IT side."),

    ("Q12", "COMPLIANCE",
     "Who at BAT signs off on outbound message templates for CASL (Canada) and CAN-SPAM (US) compliance?",
     "Anti-spam compliance is BAT's legal responsibility. Without a named owner, templates cannot ship."),

    ("Q13", "ACCEPTANCE",
     "What does '90-day post-launch success' look like in concrete numbers?",
     "Without quantitative targets, the month-3 review has no objective answer and stakeholders disagree on whether it is working."),

    ("Q14", "LAUNCH",
     "Will BAT commit to a 2-week silent pilot before outbound messaging is enabled?",
     "Silent pilot lets BD leadership review real scored outputs before any message leaves BAT's walls."),
]


def render_part(label, title, subtitle, color_hex, questions):
    # Section bar
    bar = doc.add_table(rows=1, cols=1)
    bcell = bar.rows[0].cells[0]
    bcell.width = Cm(17)
    shade_cell(bcell, color_hex)
    set_cell_borders(bcell, color=color_hex, sz="0")
    set_cell_margins(bcell, top=100, bottom=100, left=160, right=160)
    remove_default_para(bcell)
    p = cell_para(bcell, after=2)
    add_run(p, label, size=8.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    p = cell_para(bcell, after=0)
    add_run(p, title, size=13, bold=True, color=WHITE)
    if subtitle:
        p = cell_para(bcell, after=0)
        add_run(p, subtitle, size=9, italic=True, color=RGBColor(0xCF, 0xD9, 0xE8))

    add_spacer(doc, 6)

    for qid, tag, question, why in questions:
        # Q line — id + tag inline
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, qid, size=10, bold=True, color=RGBColor(*[int(color_hex[i:i+2], 16) for i in (0, 2, 4)]))
        add_run(p, "   ·   ", size=9, color=LIGHT_GREY)
        add_run(p, tag, size=8.5, bold=True, color=LIGHT_GREY)

        # Question
        add_para(doc, question, size=11, bold=True, color=NAVY,
                 after=2, line_spacing=1.3)

        # Why (one line, italic, grey)
        add_para(doc, why, size=9.5, italic=True, color=GREY,
                 after=6, line_spacing=1.35)

    add_spacer(doc, 8)


render_part("PART A", "Critical decisions",
            "Define what NORAD is — answer these first.",
            ACCENT_HEX, PART_A)

render_part("PART B", "Operational decisions",
            "Define how NORAD runs — outreach, integrations, launch.",
            AMBER_HEX, PART_B)


# ---- FOOTER ----
add_spacer(doc, 6)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.2
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p, "Prepared April 2026 by Growth Machine for BAT  ·  Project: NORAD 1",
        size=8.5, italic=True, color=LIGHT_GREY)


# ============ SAVE ============
import os
out_dir = "attached_assets"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "NORAD_DISCOVERY_QUESTIONS.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Size: {os.path.getsize(out_path)} bytes")
