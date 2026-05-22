"""Generate NORAD_BLUEPRINT.docx — stakeholder-ready Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x2A, 0x4D)
ACCENT = RGBColor(0x2E, 0x6B, 0xB8)
GREY = RGBColor(0x55, 0x5B, 0x66)
LIGHT_GREY = RGBColor(0xE8, 0xEA, 0xEE)
GOOD = RGBColor(0x1E, 0x77, 0x4F)
WARN = RGBColor(0xB6, 0x6A, 0x00)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BFC4CC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        borders.append(b)
    tc_pr.append(borders)


def add_heading(doc, text, level=1, color=NAVY, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.name = "Calibri"
    if size is None:
        size = {1: 20, 2: 15, 3: 12}.get(level, 11)
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, size=10.5, color=None, bold=False, italic=False, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_kv_table(doc, rows, col_widths=None):
    """Two-column key-value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    if col_widths is None:
        col_widths = [Cm(5), Cm(11)]
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.width = col_widths[j]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                p._element.getparent().remove(p._element)
            p = cell.add_paragraph()
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if j == 0:
                run.bold = True
                run.font.color.rgb = NAVY
                shade_cell(cell, "F2F4F8")
    return table


def add_data_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    if col_widths is None:
        total = 16
        col_widths = [Cm(total / len(headers))] * len(headers)

    # header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.width = col_widths[j]
        set_cell_borders(cell)
        shade_cell(cell, "1F2A4D")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
        p = cell.add_paragraph()
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        run.font.size = Pt(10)

    # data rows
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.width = col_widths[j]
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if i % 2 == 0:
                shade_cell(cell, "F7F8FB")
            for p in cell.paragraphs:
                p._element.getparent().remove(p._element)
            p = cell.add_paragraph()
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
    return table


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1F2A4D')
    pBdr.append(bottom)
    pPr.append(pBdr)


def page_break(doc):
    doc.add_page_break()


# ========== BUILD DOCUMENT ==========

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# === COVER ===
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(120)
run = p.add_run("PROJECT NORAD 1")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.font.name = "Calibri"

p = doc.add_paragraph()
run = p.add_run("BD Intelligence & Deal-Sourcing Engine")
run.font.size = Pt(18)
run.font.color.rgb = ACCENT
run.font.name = "Calibri"

p = doc.add_paragraph()
run = p.add_run("Project Blueprint  ·  v1 Build Plan")
run.italic = True
run.font.size = Pt(13)
run.font.color.rgb = GREY
run.font.name = "Calibri"

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
run = p.add_run("Prepared by Growth Machine for BAT  ·  April 2026")
run.font.size = Pt(11)
run.font.color.rgb = GREY
run.font.name = "Calibri"

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("First tenant: BAT  ·  Designed niche-agnostic")
run.font.size = Pt(11)
run.font.color.rgb = GREY
run.font.name = "Calibri"

page_break(doc)

# === 1. EXECUTIVE SUMMARY ===
add_heading(doc, "1.  Executive Summary", level=1)
add_hr(doc)

add_para(doc,
    "NORAD is an automated business-development intelligence engine. Twice a day it pulls public "
    "signals from regulatory databases, public filings, patents, and trade press across Canada and the US, "
    "links every signal to the right company, scores each company against a configurable rubric, and "
    "produces a small ranked list of acquisition or partnership targets ready for human review.")

add_para(doc,
    "BAT is the first tenant. The system is designed multi-tenant from day one — the same engine can "
    "serve a fintech VC, a pharma corp-dev team, or a healthcare scout with configuration changes only, "
    "no re-architecture.")

add_para(doc, "Hard rules:", bold=True)
add_bullet(doc, "Nothing sends automatically. Every outbound message is approved by a human.")
add_bullet(doc, "Every signal links back to its public source.")
add_bullet(doc, "Every score has an auditable reason.")
add_bullet(doc, "Multi-tenant from day one.")

# === 2. THE BAT USE CASE ===
add_heading(doc, "2.  The BAT Use Case", level=1)
add_hr(doc)

add_kv_table(doc, [
    ("Problem",
     "BD analysts manually scan regulatory filings, trade press, patents and distributor news to "
     "find emerging convenience-store and pharmacy-retail brands. Slow, inconsistent, and most signals are missed."),
    ("Solution",
     "Twice-daily automated discovery → cleaning → entity resolution → scoring → outreach drafting, "
     "with a human-in-the-loop approval queue before anything leaves the system."),
    ("Users",
     "BAT's BD analysts and BD leadership; Marketing approves outreach copy."),
    ("Outputs",
     "(1) Live Discover feed of every signal. (2) Ranked Opportunities list with score + justification + red flags. "
     "(3) One-page deal-card PDF per opportunity. (4) Draft outreach pack ready for analyst finalisation. "
     "(5) Review Queue for human approval."),
    ("Scope",
     "v1 covers Canada + US, convenience-store and pharmacy-retail-adjacent CPG (functional beverages, OTC, supplements, snacks)."),
])

page_break(doc)

# === 3. CORE TOOLS ===
add_heading(doc, "3.  Core Tools — What, Why, and Cost", level=1)
add_hr(doc)

add_para(doc,
    "Pricing below is indicative as of April 2026, pulled from each vendor's published pricing page. "
    "Final cost depends on negotiated volume tier. v1 estimates assume one tenant (BAT) and modest discovery "
    "volume of approximately 5,000–20,000 external API calls per day.",
    italic=True, color=GREY, size=9.5)

# 3.1 Frontend
add_heading(doc, "3.1  Frontend (what BAT analysts click on)", level=2)
add_data_table(doc,
    headers=["Tool", "Role", "Why this", "Indicative cost (v1)"],
    rows=[
        ("Next.js 15", "Frontend framework",
         "Server-side rendering for shareable deal-card URLs; large hiring pool; best ecosystem.",
         "Included with Vercel"),
        ("TypeScript", "Language end-to-end",
         "Same language as backend → schemas shared, no client/server drift.",
         "Free"),
        ("Tailwind + shadcn/ui", "Styling + components",
         "Components copied into repo (no library lock-in); industry default.",
         "Free"),
        ("Clerk", "Auth, multi-tenant orgs, RBAC",
         "Saves 4+ weeks of identity work; orgs map directly to tenants.",
         "$0/mo (free under 50K MAU)"),
        ("tRPC", "Internal typed API",
         "End-to-end types between frontend and backend, zero codegen.",
         "Free"),
    ],
    col_widths=[Cm(2.8), Cm(3.2), Cm(7.2), Cm(2.8)],
)

# 3.2 Backend
add_heading(doc, "3.2  Backend (the engine)", level=2)
add_data_table(doc,
    headers=["Tool", "Role", "Why this", "Indicative cost (v1)"],
    rows=[
        ("Node.js 22 LTS + TypeScript", "Runtime + language",
         "One language top-to-bottom; massive ecosystem.", "Free"),
        ("Fastify", "HTTP server",
         "Faster, smaller, schema-first vs Express.", "Free"),
        ("Drizzle ORM", "Database access",
         "TypeScript-native, no runtime overhead, predictable migrations.", "Free"),
        ("Inngest", "Durable workflow engine",
         "Step-level retries, replay, cron + event triggers, TS-native. The right tool "
         "for the fetch → clean → resolve → score → notify pipeline.",
         "$0/mo free tier (50K runs); $20/mo Basic later"),
    ],
    col_widths=[Cm(2.8), Cm(3.2), Cm(7.2), Cm(2.8)],
)

# 3.3 AI / Agents
add_heading(doc, "3.3  AI / Agent Layer (used at only 2 sites)", level=2)
add_data_table(doc,
    headers=["Tool", "Role", "Why this", "Indicative cost (v1)"],
    rows=[
        ("Mastra", "TypeScript agent framework",
         "Built on Vercel AI SDK; production-ready agents with memory + workflows; same monorepo.",
         "Free (open source)"),
        ("OpenAI GPT-4o-mini", "The only LLM — unstructured-doc parser + score justification writer",
         "Best $/quality in its tier. We don't need a frontier model for these two narrow jobs.",
         "$0.15 / 1M input · $0.60 / 1M output → ~$30–100/mo"),
        ("OpenAI text-embedding-3-small", "Embeddings for semantic name matching",
         "Best $/quality; same vendor → one API key, one billing relationship.",
         "$0.02 / 1M tokens → ~$5/mo"),
    ],
    col_widths=[Cm(3.2), Cm(3.6), Cm(6.4), Cm(2.8)],
)

# 3.4 Entity Resolution
add_heading(doc, "3.4  Entity Resolution (the hardest data problem)", level=2)
add_data_table(doc,
    headers=["Tool", "Role", "Why this", "Indicative cost (v1)"],
    rows=[
        ("Splink (Python service)", "Probabilistic record linkage — links 'Northern Wellness Inc', "
         "'Northern Wellness LLC', 'NorthernWell' to one company",
         "UK Ministry of Justice maintains it; MIT licence; scales to tens of millions of records.",
         "Free; ~$10–30/mo Fly.io compute"),
        ("pgvector + pgvectorscale", "Vector similarity inside Postgres",
         "11.4× throughput vs vanilla pgvector; lives inside Postgres → no second DB to operate.",
         "Free (Postgres extension)"),
    ],
    col_widths=[Cm(3.2), Cm(4.0), Cm(6.0), Cm(2.8)],
)

page_break(doc)

# 3.5 Data stores
add_heading(doc, "3.5  Data Stores", level=2)
add_data_table(doc,
    headers=["Tool", "Role", "Why this", "Indicative cost (v1)"],
    rows=[
        ("Postgres 16 on Supabase", "Primary OLTP DB — entities, events, scores, configs, users",
         "Managed Postgres + RLS for tenant isolation + storage; scales cleanly.",
         "$25/mo Pro (8 GB included)"),
        ("Redis 7 on Upstash", "Cache, queues (BullMQ), rate-limit counters",
         "Pay-per-use, no idle cost, serverless. Avoids managing a Redis box.",
         "$5–20/mo PAYG"),
        ("Cloudflare R2", "Object storage for raw payloads + generated PDFs",
         "Zero egress fees (huge over time vs S3); $0.015/GB/mo storage; 10 GB free.",
         "~$5/mo"),
    ],
    col_widths=[Cm(3.2), Cm(3.8), Cm(6.2), Cm(2.8)],
)

# 3.6 Bought source connectors
add_heading(doc, "3.6  Bought Source Connectors (the four 'outside-world' tools)", level=2)
add_data_table(doc,
    headers=["Tool", "Role", "Why this", "Indicative cost (v1)"],
    rows=[
        ("Exa", "Semantic web search — 'find URLs about emerging Canadian functional-beverage launches in last 7 days'",
         "AI-native search; understands meaning not just keywords. Replaces months of building our own crawler.",
         "$7 per 1K searches (1K free/mo) → ~$10–50/mo"),
        ("Firecrawl", "URL → clean markdown. Handles JS-rendered pages, login walls, dynamic content.",
         "Without it we'd parse raw HTML soup. Dominant tool for LLM-ready extraction.",
         "$83/mo Standard (100K pages/mo)"),
        ("Bright Data SERP API", "High-volume Google searches — e.g. 'Northern Wellness Inc news 2026'",
         "Cheapest at scale; failed requests not billed. Different job from Exa: returns Google's actual ranking.",
         "$1.80–$3.00 per 1K req → ~$10–50/mo (Micro $10 commit)"),
        ("Diffbot Knowledge Graph", "Company entity enrichment — name in, structured facts out (HQ, size, founders, web, products)",
         "Pre-built, structured, instantly fills metadata we'd otherwise scrape.",
         "$299/mo Startup (250K credits/mo)"),
    ],
    col_widths=[Cm(3.2), Cm(4.6), Cm(5.4), Cm(2.8)],
)

# 3.7 Built-in connectors
add_heading(doc, "3.7  Built-in Source Connectors (free, public APIs)", level=2)

add_para(doc, "Not products we pay for — public APIs / bulk downloads we wrap in our own connector code. "
    "These become the moat of the system because no one else bothers to build them.", italic=True, color=GREY, size=10)

add_data_table(doc,
    headers=["Source", "Country", "Data", "Cost"],
    rows=[
        ("Health Canada NHP", "CA", "Natural Health Product approvals, licence changes, Notices of Non-Compliance", "Free"),
        ("OpenFDA", "US", "GRAS, NDI, OTC monograph status, Warning Letters, Import Refusals", "Free"),
        ("SEDAR+", "CA", "Public-company filings", "Free"),
        ("SEC EDGAR", "US", "S-1, 10-K, 8-K, M&A disclosures", "Free"),
        ("USPTO + CIPO", "US + CA", "Patent filings (formula tech, packaging IP, brand IP)", "Free"),
    ],
    col_widths=[Cm(3.6), Cm(2.0), Cm(8.0), Cm(2.4)],
)

page_break(doc)

# 3.8 Outbound
add_heading(doc, "3.8  Outbound (what leaves the system)", level=2)
add_data_table(doc,
    headers=["Tool", "Role", "Why this", "Indicative cost (v1)"],
    rows=[
        ("SendGrid", "Email send for human-approved outreach",
         "Mature, reliable deliverability, generous quota.", "~$20/mo Essentials (50K emails/mo)"),
        ("HubSpot", "CRM sync (one-way push v1; bidirectional v1.5)",
         "Enterprise standard; free CRM tier covers 1M contacts; integration itself is free.",
         "$0/mo (uses BAT's existing HubSpot)"),
        ("React-PDF", "Generates the one-page deal card",
         "Pure React, no headless browser to operate, deterministic output.", "Free"),
    ],
    col_widths=[Cm(3.0), Cm(3.6), Cm(6.6), Cm(2.8)],
)

# 3.9 Observability + hosting
add_heading(doc, "3.9  Observability, Secrets, and Hosting", level=2)
add_data_table(doc,
    headers=["Tool", "Role", "Why this", "Indicative cost (v1)"],
    rows=[
        ("Langfuse", "LLM tracing — every call logged with cost, latency, prompt version",
         "Non-negotiable from day 1. Without it, agent regressions go silent.",
         "$29/mo Core (100K units/mo, unlimited users)"),
        ("Doppler", "Secrets vault",
         "Free for 3 users; clean SDK; no rotating env files in repo.",
         "$0/mo (free tier)"),
        ("Vercel", "Frontend hosting (Next.js)",
         "Built for Next.js, edge network, painless deploys.",
         "$20/mo Pro (1 seat, $20 usage credit)"),
        ("Fly.io", "Backend container hosting (api, worker, resolver-py)",
         "Pay-per-second, runs containers globally, no AWS-tax learning curve.",
         "~$30–60/mo"),
    ],
    col_widths=[Cm(2.4), Cm(4.0), Cm(6.8), Cm(2.8)],
)

# === 4. COST SUMMARY ===
add_heading(doc, "4.  Indicative Monthly Cost Summary (v1, one tenant)", level=1)
add_hr(doc)

add_data_table(doc,
    headers=["Bucket", "Monthly cost (low → high)"],
    rows=[
        ("AI (LLM + embeddings)", "$35 – $105"),
        ("Bought source connectors (Exa + Firecrawl + Bright Data + Diffbot)", "$402 – $482"),
        ("Data stores (Supabase + Upstash + R2)", "$35 – $50"),
        ("Outbound (SendGrid; HubSpot integration is free)", "$20"),
        ("Observability + secrets (Langfuse + Doppler)", "$29"),
        ("Hosting (Vercel + Fly.io)", "$50 – $80"),
        ("Auth (Clerk free tier under 50K MAU)", "$0"),
        ("Workflow engine (Inngest free tier)", "$0"),
        ("Total at v1 scale", "≈ $570 – $770 / month"),
    ],
    col_widths=[Cm(11), Cm(5)],
)

add_para(doc, "")
add_para(doc, "At second-tenant onboarding: add ~$50–150/mo per tenant (mostly variable Diffbot + Exa usage). "
    "Fixed costs (hosting, auth, observability) are amortised across tenants.", color=GREY, size=10)

page_break(doc)

# === 5. PERPLEXITY SONAR ===
add_heading(doc, "5.  Why Perplexity Sonar Is Not in v1 (and When It Earns Its Slot)", level=1)
add_hr(doc)

add_para(doc, "What Sonar does.", bold=True)
add_para(doc,
    "Sonar is an 'ask a natural-language question, get a grounded answer with citations' API. "
    "Pricing: $1/$1 per 1M tokens (Sonar base), $3/$15 per 1M tokens (Sonar Pro).")

add_para(doc, "Why it is not in v1.", bold=True)
add_para(doc, "We already produce the same outcome by composing tools we already have:")
add_para(doc, "    Exa (find URLs)  →  Firecrawl (extract clean text)  →  GPT-4o-mini (synthesise + cite)",
    italic=True, color=ACCENT, indent=0.5)
add_para(doc, "That stack gives us:")
add_bullet(doc, "Full control over the prompt and output schema (Sonar's output is opaque).")
add_bullet(doc, "Lower per-call cost at our volume (Sonar Pro output $15 / 1M vs GPT-4o-mini $0.60 / 1M).")
add_bullet(doc, "Visibility into every retrieved source (we own the URL list; Sonar surfaces citations but doesn't let us audit retrieval).")

add_para(doc, "When Sonar earns a slot.", bold=True)
add_para(doc,
    "Phase 2, when we add a Natural Language Search (NLS) feature on the Discover page — for example "
    "'show me emerging functional beverage brands with regulatory tailwind in Quebec' — and we want a "
    "synthesised, cited answer back without writing the orchestration ourselves. Sonar saves roughly "
    "2–4 weeks of agent-glue code for that one feature.")

add_para(doc, "Verdict: Phase 2 add-on, not v1. Estimated added cost: ~$50–200/mo depending on NLS adoption.",
    bold=True, color=NAVY)

# === 6. SWEET SPOT ===
add_heading(doc, "6.  The Sweet Spot — Beyond BAT", level=1)
add_hr(doc)

add_para(doc, "The system is designed around one principle: BAT is the first tenant, not the only one. "
    "Adding more sources, more tools, and more niches later is configuration and connector work, not a re-architecture.",
    bold=True)

add_heading(doc, "6.1  Three-layer architecture", level=2)
add_data_table(doc,
    headers=["Layer", "What lives here", "Change frequency"],
    rows=[
        ("Tenant config (JSON in DB)", "Sources, rubric, red flags, voice, approval workflow", "Daily, by users in the UI"),
        ("Domain adapters (code modules)", "Source connectors, scrapers, AI providers, CRMs", "Monthly, by developers"),
        ("Engine core", "Pipeline, resolver, scoring math, workflows", "Quarterly, rarely touched"),
    ],
    col_widths=[Cm(4.5), Cm(8.0), Cm(3.5)],
)

add_heading(doc, "6.2  Adding a new source", level=2)
add_para(doc,
    "Every source implements one interface — fetchSince(date) → RawEvent[]. One new file, 50–200 lines, "
    "no engine changes. The source then appears automatically in the Source Manager UI for tenants to enable.")

add_heading(doc, "6.3  Adding a new tool", level=2)
add_para(doc,
    "The system uses an adapter pattern at every external boundary. Swap one config or one adapter file. "
    "The agents and pipelines do not know or care which model, scraper, or CRM is behind them.")

add_heading(doc, "6.4  Adding a new niche / tenant", level=2)
add_para(doc, "None of the niche-specific logic is in code. It all lives in per-tenant configuration:")
add_data_table(doc,
    headers=["Niche-specific concern", "Where it lives"],
    rows=[
        ("Which sources to monitor", "tenant_configs.sources (JSON)"),
        ("Rubric weights & metric definitions", "tenant_configs.rubric (JSON)"),
        ("Red-flag rules", "tenant_configs.red_flags (JSON)"),
        ("Tone of voice for outreach", "tenant_configs.voice (JSON)"),
        ("Approval workflow chain", "tenant_configs.approval_workflow (JSON)"),
    ],
    col_widths=[Cm(8), Cm(8)],
)
add_para(doc,
    "Onboarding a new tenant = a few hours of config work in the UI. Zero engineering work, "
    "assuming the source library already covers their niche.", bold=True, color=NAVY)

add_heading(doc, "6.5  Future tools and sources we can plug in later", level=2)
add_data_table(doc,
    headers=["Category", "Future additions (not v1)"],
    rows=[
        ("Social signals", "LinkedIn (exec moves), Reddit (community discovery), TikTok / Instagram (DTC discovery)"),
        ("Discovery breadth", "Trade-show attendee lists, conference speaker lists, job postings (Greenhouse / Lever)"),
        ("Web intelligence", "SimilarWeb (traffic), Wappalyzer (tech stack)"),
        ("AI capability", "Perplexity Sonar (NLS), Claude / GPT-5 routing, vision models for label scanning"),
        ("CRM", "Salesforce, Pipedrive (alongside HubSpot)"),
        ("Distributor data", "UNFI portal, KeHE portal, McKesson (often gated → partnership required)"),
        ("Analyst overlays", "Crunchbase, PitchBook (M&A history), Owler (private-co revenue estimates)"),
    ],
    col_widths=[Cm(4.5), Cm(11.5)],
)
add_para(doc, "Each one slots in via the same Connector interface. No re-architecture.",
    italic=True, color=GREY)

page_break(doc)

# === 7. WHAT'S DEFERRED ===
add_heading(doc, "7.  What Is Deferred (and Why That Is a Choice, Not an Oversight)", level=1)
add_hr(doc)

add_data_table(doc,
    headers=["Deferred to", "Why"],
    rows=[
        ("LinkedIn / Reddit / Trade-show data → Phase 2",
         "Legally tricky and/or requires Bright Data residential or Phantombuster — bigger engineering and compliance lift. "
         "Source coverage is good enough for v1 with regulatory + filings + patents + trade press."),
        ("Perplexity Sonar (NLS) → Phase 2",
         "We already do its job by composing Exa + Firecrawl + GPT-4o-mini at lower cost. "
         "Add Sonar only when we ship the NL Search UX."),
        ("Two-way HubSpot sync → v1.5",
         "One-way push (NORAD → HubSpot) ships in v1. Bidirectional sync (read existing CRM contacts → "
         "avoid double-outreach) is finicky and a known follow-up."),
        ("Meilisearch → optional v1 add",
         "Postgres full-text search works at v1 volumes. Add Meilisearch when Discover page latency demands it."),
        ("Terraform full IaC → v1.5",
         "Manual provisioning of Vercel + Supabase + Upstash + Fly.io is fine for one tenant; "
         "Terraform pays off at 3+ environments."),
    ],
    col_widths=[Cm(5.5), Cm(10.5)],
)

# === 8. CONFIDENCE ===
add_heading(doc, "8.  Honest Confidence Score", level=1)
add_hr(doc)

add_data_table(doc,
    headers=["Area", "Confidence", "Note"],
    rows=[
        ("Architecture & extensibility", "9 / 10", "Proven 3-layer pattern; niche-agnostic by design."),
        ("Stack choices", "8.5 / 10", "All boring, battle-tested tools with clean escape hatches."),
        ("Source coverage for BAT v1", "7 / 10", "Strong on regulatory / filings / patents; missing social until Phase 2."),
        ("Time-to-value", "7.5 / 10", "Entity resolution + scoring tuning needs 4–6 weeks to sharpen post-launch."),
        ("Future expansion (more niches/tools)", "9 / 10", "This is the design's strongest property."),
        ("Overall", "≈ 8 / 10", "Strong v1 plan with two known soft spots that are explicit, not hidden."),
    ],
    col_widths=[Cm(5.5), Cm(2.5), Cm(8.0)],
)

add_para(doc, "")
add_para(doc, "To raise confidence to 9.5 / 10 before build starts:", bold=True)
add_bullet(doc, "Confirm BAT's weekly target volume of qualified opportunities (5–10/wk vs 30+/wk).")
add_bullet(doc, "Pre-load 50 BAT-curated example companies as gold-standard scoring data.")
add_bullet(doc, "Plan a 2-week silent pilot (engine runs, no outreach goes out, BAT validates ranked lists).")
add_bullet(doc, "Decide whether to pull LinkedIn signals into v1 (legal + cost call).")

# === FOOTER ===
add_para(doc, "")
add_hr(doc)
add_para(doc, "Prepared April 2026 by Growth Machine for BAT  ·  Project: NORAD 1  ·  All vendor pricing pulled from official pricing pages; confirm current rates at vendor sign-up.",
    italic=True, color=GREY, size=9)

# Save
import os
out_dir = "attached_assets"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "NORAD_BLUEPRINT.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Size: {os.path.getsize(out_path)} bytes")
